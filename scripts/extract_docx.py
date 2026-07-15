"""
Generic .docx extractor — reusable for any PreOne document.
Extracts:
  - paragraphs  → {name}_paragraphs.txt  (index | style | text)
  - tables      → {name}_tables.json     (full structure)
  - summary     → {name}_tables_summary.txt
"""
import sys, os, json
from docx import Document

def extract(path, out_name):
    doc = Document(path)
    base = f"/home/z/my-project/extracted/{out_name}"
    os.makedirs(os.path.dirname(base), exist_ok=True)

    # ----- paragraphs -----
    with open(f"{base}_paragraphs.txt", "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            style = p.style.name if p.style is not None else "Normal"
            f.write(f"[{i:05d}] [{style}] {txt}\n")

    # ----- tables -----
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append({"index": ti, "rows": len(rows),
                       "cols": len(rows[0]) if rows else 0, "data": rows})
    with open(f"{base}_tables.json", "w", encoding="utf-8") as f:
        json.dump(tables, f, ensure_ascii=False, indent=2)

    # ----- tables summary -----
    with open(f"{base}_tables_summary.txt", "w", encoding="utf-8") as f:
        f.write(f"# {out_name} — {len(tables)} tables\n\n")
        for t in tables:
            f.write(f"## Table #{t['index']}  rows={t['rows']} cols={t['cols']}\n")
            if t['data']:
                hdr = t['data'][0]
                f.write("Header: " + " | ".join(hdr) + "\n")
                for r in t['data'][1:4]:
                    preview = " | ".join(c[:60] for c in r)
                    f.write("  Row: " + preview + "\n")
            f.write("\n")

    print(f"✓ {out_name}: {len(doc.paragraphs)} paragraphs, {len(tables)} tables")
    print(f"  → {base}_paragraphs.txt")
    print(f"  → {base}_tables.json")
    print(f"  → {base}_tables_summary.txt")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python extract_docx.py <path> <out_name>")
        sys.exit(1)
    extract(sys.argv[1], sys.argv[2])
