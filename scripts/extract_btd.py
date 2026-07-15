"""
Extract PreOne_Backend_Technical_Design_v1.0.docx
"""
from docx import Document
from pathlib import Path
import json

SRC = "/home/z/my-project/upload/PreOne_Backend_Technical_Design_v1.0.docx"
OUT_DIR = Path("/home/z/my-project/extracted")
OUT_DIR.mkdir(parents=True, exist_ok=True)

doc = Document(SRC)

# ---------- Paragraphs ----------
para_lines = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else ""
    if text:
        para_lines.append(f"[{i:04d}] ({style}) {text}")

(OUT_DIR / "btd_paragraphs.txt").write_text(
    "\n".join(para_lines), encoding="utf-8"
)
print(f"Backend Tech Design Paragraphs: {len(para_lines)} non-empty / {len(doc.paragraphs)} total")

# ---------- Tables ----------
tables_out = []
table_summary = []
for ti, table in enumerate(doc.tables):
    rows = []
    for r in table.rows:
        cells = [c.text.strip().replace("\n", " | ") for c in r.cells]
        rows.append(cells)
    tables_out.append({"index": ti, "rows": rows})
    header = rows[0] if rows else []
    table_summary.append({
        "index": ti,
        "row_count": len(rows),
        "col_count": len(header),
        "header_preview": header[:6],
    })

(OUT_DIR / "btd_tables.json").write_text(
    json.dumps(tables_out, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(f"Backend Tech Design Tables: {len(doc.tables)}")

with open(OUT_DIR / "btd_tables_summary.txt", "w", encoding="utf-8") as f:
    for s in table_summary:
        f.write(f"--- Table #{s['index']} ---\n")
        f.write(f"Rows: {s['row_count']}  Cols: {s['col_count']}\n")
        f.write(f"Header: {s['header_preview']}\n\n")

print("\n=== All BTD tables preview ===")
for s in table_summary:
    print(f"#{s['index']:03d}  rows={s['row_count']:4d}  cols={s['col_count']}  header={s['header_preview']}")
